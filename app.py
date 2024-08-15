import os
import requests
from dotenv import load_dotenv
import docx
from docx import Document
import re
import shutil
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from translate_prompt import extract_text_with_structure, create_translation_prompt, create_error_report_prompt, create_improvement_prompt, create_natural_translation_prompt, create_error_free_translation_prompt
import firebase_admin
from firebase_admin import credentials, storage
from datetime import timedelta
from flask_cors import CORS  # Import CORS

# Initialize Firebase Admin SDK
cred = credentials.Certificate('firebase-adminsdk.json')
firebase_admin.initialize_app(cred, {
    'storageBucket': 'tomato-4836e.appspot.com'
})

# Load environment variables from the .env file
load_dotenv()
# Accessing the environment variables
flask_env = os.getenv("FLASK_ENV")
secret_key = os.getenv("SECRET_KEY")
# Retrieve the API key from the environment variables
api_key = os.getenv("GEMINI_API_KEY")

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes and origins by default
app.config['ENV'] = flask_env
app.config['SECRET_KEY'] = secret_key
# Define the allowed file extensions
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Chunking Function
def chunk_docx_by_paragraphs(input_file, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    else:
        shutil.rmtree(output_dir)  # Clear the directory if it exists
        os.makedirs(output_dir)

    doc = Document(input_file)
    chunk_index = 0
    output_files = []

    for paragraph in doc.paragraphs:
        new_doc = Document()
        new_doc.add_paragraph(paragraph.text)

        output_file = os.path.join(output_dir, f"chunk_{chunk_index}.docx")
        new_doc.save(output_file)
        output_files.append(output_file)
        print(f"Saved chunk {chunk_index} successfully.")
        chunk_index += 1

    return output_files


def upload_to_firebase_storage(local_file_path, storage_path):
    """Uploads a file to Firebase Storage and returns its public URL."""
    bucket = storage.bucket()
    blob = bucket.blob(storage_path)
    blob.upload_from_filename(local_file_path)
    print(f"File {local_file_path} uploaded to {storage_path}.")
    # Make the file publicly accessible and get the public URL
    blob.make_public()
    return blob.public_url

def get_signed_url(storage_path):
    """Generates a signed download URL for a file in Firebase Storage."""
    bucket = storage.bucket()
    blob = bucket.blob(storage_path)
    return blob.generate_signed_url(timedelta(seconds=300))

# Translation Function
def translate_texts(source_lang, target_lang, country, writer):
    token_counts = []
    total_tokens_used = 0
    source_folder = "original_chunks"
    final_translation_file = "final_translation.docx"
    final_doc = Document()

    files = [f for f in os.listdir(source_folder) if f.startswith("chunk_") and f.endswith(".docx")]
    files.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))

    for i, filename in enumerate(files):
        source_docx_file = os.path.join(source_folder, filename)
        print(f"Processing {source_docx_file}")

        elements = extract_text_with_structure(source_docx_file)
        text_content = "".join(element[1] for element in elements).strip()
        if not text_content:
            print(f"Original content for {source_docx_file} contains only line breaks or is empty.")
            append_original_to_docx(final_doc, source_docx_file)
            continue

        translation_prompt = create_translation_prompt(source_lang, target_lang, country, elements)

        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={api_key}"
        headers = {
            "Content-Type": "application/json"
        }

        def call_gemini_api(prompt):
            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "text": prompt
                            }
                        ]
                    }
                ]
            }
            response = requests.post(url, headers=headers, json=payload)
            if response.status_code == 200:
                response_data = response.json()
                print("Response JSON:", response_data)
                token_count = response_data.get('usageMetadata', {}).get('totalTokenCount', 0)
                return response_data, token_count
            else:
                print(f"Error: {response.status_code} - {response.text}")
                return None, 0

        def process_translation(prompt):
            translation_response, tokens = call_gemini_api(prompt)
            token_counts.append(tokens)

            if translation_response and 'candidates' in translation_response and 'content' in translation_response['candidates'][0]:
                return translation_response, tokens
            else:
                print("Retrying due to missing 'content' in response...")
                translation_response, tokens = call_gemini_api(prompt)
                token_counts.append(tokens)
                if translation_response and 'candidates' in translation_response and 'content' in translation_response['candidates'][0]:
                    return translation_response, tokens
                else:
                    print("Translation failed after retry.")
                    return None, tokens

        translation_response, tokens = process_translation(translation_prompt)
        if translation_response:
            translated_text = translation_response['candidates'][0]['content']['parts'][0]['text']
            translated_text = re.sub(r'<[^>]+>', '', translated_text).strip()

            error_report_prompt = create_error_report_prompt(source_lang, target_lang, country, elements, translated_text)
            error_report_response, tokens = process_translation(error_report_prompt)
            if error_report_response:
                error_report_text = error_report_response['candidates'][0]['content']['parts'][0]['text']

                improvement_prompt = create_improvement_prompt(source_lang, target_lang, country, translated_text, error_report_text)
                improvement_response, tokens = process_translation(improvement_prompt)
                if improvement_response:
                    improvement_text = improvement_response['candidates'][0]['content']['parts'][0]['text']
                    improvement_text = re.sub(r'<[^>]+>', '', improvement_text).strip()

                    natural_translation_prompt = create_natural_translation_prompt(target_lang, country, improvement_text, writer)
                    natural_translation_response, tokens = process_translation(natural_translation_prompt)
                    if natural_translation_response:
                        natural_translation_text = natural_translation_response['candidates'][0]['content']['parts'][0]['text'].replace('\n\n','')
                        natural_translation_text = re.sub(r'<[^>]+>', '', natural_translation_text).strip()

                        error_free_translation_prompt = create_error_free_translation_prompt(target_lang, country, natural_translation_text, writer)
                        error_free_translation_response, tokens = process_translation(error_free_translation_prompt)
                        if error_free_translation_response:
                            error_free_translation_text = error_free_translation_response['candidates'][0]['content']['parts'][0]['text'].replace('\n\n','')
                            error_free_translation_text = re.sub(r'<[^>]+>', '', error_free_translation_text).strip()

                            append_translation_to_docx(final_doc, error_free_translation_text, elements)
                            continue

        print(f"Appending original content for {source_docx_file} due to translation failure or missing content.")
        append_original_to_docx(final_doc, source_docx_file)

    total_tokens_used = sum(token_counts)
    print(f"Total tokens used: {total_tokens_used}")

    final_doc.save(final_translation_file)
    print(f"All translations saved to {final_translation_file}")

    # Upload the file to Firebase Storage
    public_url = upload_to_firebase_storage(final_translation_file, final_translation_file)
    # Alternatively, use a signed URL:
    # signed_url = get_signed_url(final_translation_file)

    return final_translation_file, total_tokens_used, public_url  # or return signed_url

# Flask Route for Translation API
@app.route('/translate', methods=['POST'])
def translate():
    # Check if the post request has the file part
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No file selected for uploading"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        input_file_path = os.path.join(os.getcwd(), filename)
        file.save(input_file_path)
        print(f"File {filename} uploaded successfully.")

        # Process the file for chunking and translation
        source_lang = request.form.get('source_lang', '')
        target_lang = request.form.get('target_lang', '')
        country = request.form.get('country', '')
        writer = request.form.get('writer', '')

        output_dir = 'original_chunks'

        # Step 1: Chunk the document by paragraphs, where each chunk is a single paragraph
        chunk_docx_by_paragraphs(input_file_path, output_dir)

        # Step 2: Translate the chunked files
        final_translation_file, total_tokens_used, download_url = translate_texts(source_lang, target_lang, country, writer)

        return jsonify({
            'status': 'success',
            'final_translation_file': final_translation_file,
            'total_tokens_used': total_tokens_used,
            'download_url': download_url  # Include the download URL in the response
        })

    return jsonify({"error": "File not allowed"}), 400

# Helper Functions
def save_text_to_docx(source_docx_file, text, elements, destination_docx_file, is_plain_text=False):
    cleaned_text = text.replace('\n\n', '')
    paragraphs = cleaned_text.split('\n')
    doc = Document()

    if is_plain_text:
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    else:
        for element, para in zip(elements, paragraphs):
            if element[0] == 'paragraph':
                new_para = doc.add_paragraph(style=element[2])
                new_run = new_para.add_run(para)
                copy_run_formatting(new_run, element[3][0] if element[3] else None)
            elif element[0] == 'cell':
                new_table = doc.add_table(rows=1, cols=1)
                cell = new_table.cell(0, 0)
                cell.text = para
                for run in element[3]:
                    new_run = cell.paragraphs[0].add_run(run.text)
                    copy_run_formatting(new_run, run)

    doc.save(destination_docx_file)
    print(f"Text saved to {destination_docx_file}")

def copy_run_formatting(new_run, original_run):
    if original_run:
        new_run.bold = original_run.bold
        new_run.italic = original_run.italic
        new_run.underline = original_run.underline
        new_run.font.size = original_run.font.size
        new_run.font.color.rgb = original_run.font.color.rgb

def append_translation_to_docx(merged_doc, translated_text, elements):
    translated_paragraphs = translated_text.split('\n')
    for element, translated_para in zip(elements, translated_paragraphs):
        if element[0] == 'paragraph':
            new_para = merged_doc.add_paragraph(style=element[2])
            new_run = new_para.add_run(translated_para)
            copy_run_formatting(new_run, element[3][0] if element[3] else None)
        elif element[0] == 'cell':
            new_table = merged_doc.add_table(rows=1, cols=1)
            cell = new_table.cell(0, 0)
            cell.text = translated_para
            for run in element[3]:
                new_run = cell.paragraphs[0].add_run(run.text)
                copy_run_formatting(new_run, run)

def append_original_to_docx(merged_doc, source_docx_file):
    original_doc = Document(source_docx_file)
    for paragraph in original_doc.paragraphs:
        new_para = merged_doc.add_paragraph(style=paragraph.style)
        for run in paragraph.runs:
            new_run = new_para.add_run(run.text)
            copy_run_formatting(new_run, run)

def append_doc_to_another(doc, doc_to_append):
    for element in iter_block_items(doc_to_append):
        if isinstance(element, docx.text.paragraph.Paragraph):
            new_para = doc.add_paragraph(style=element.style)
            for run in element.runs:
                new_run = new_para.add_run(run.text)
                copy_run_formatting(new_run, run)
        elif isinstance(element, docx.table.Table):
            new_table = doc.add_table(rows=0, cols=len(element.columns))
            for row in element.rows:
                new_row = new_table.add_row()
                for i, cell in enumerate(row.cells):
                    new_cell = new_row.cells[i]
                    for paragraph in cell.paragraphs:
                        new_para = new_cell.add_paragraph(style=paragraph.style)
                        for run in paragraph.runs:
                            new_run = new_para.add_run(run.text)
                            copy_run_formatting(new_run, run)

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._element
    else:
        raise ValueError("Unknown parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.CT_Tbl):
            yield docx.table.Table(child, parent)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=8080, debug=True)
