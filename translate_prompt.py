import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tc

def extract_text_with_structure(docx_file):
    doc = docx.Document(docx_file)
    elements = []

    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            elements.append(('paragraph', get_full_text_from_paragraph(block), block.style, block.runs))
        elif isinstance(block, docx.table._Cell):
            for paragraph in block.paragraphs:
                elements.append(('cell', get_full_text_from_paragraph(paragraph), paragraph.style, paragraph.runs))

    return elements

def iter_block_items(parent):
    """
    Yield each block item from a parent, allowing for paragraphs and table cells.
    """
    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tc):
            for cell in docx.table._Cell(child, parent).tables:
                for row in cell.rows:
                    for table_cell in row.cells:
                        yield table_cell

def get_full_text_from_paragraph(paragraph):
    """
    Concatenate text from all runs in a paragraph to maintain formatting and avoid extra line breaks.
    """
    full_text = ""
    for run in paragraph.runs:
        full_text += run.text
    return full_text

# Function to return translation guidelines
def get_translation_guidelines(target_lang, country):
    return f"""
1. **Style Consistency**: Ensure the style matches the original text. Consider the cultural and stylistic nuances of {target_lang} as spoken in {country}. For example, a news article should read like a news article in {target_lang}.
2. **Avoid Errors**: Prevent mistranslations, omissions, or untranslated text.
3. **Accuracy Above All**: Accuracy is paramount. Your translation must convey the correct meaning, taking precedence over all other considerations.
"""


def create_translation_prompt(source_lang, target_lang, country, elements):
    translation_prompt = f"""
    Please translate the following source text from {source_lang} to {target_lang}.

    ### Translation Guidelines:
    {get_translation_guidelines(target_lang, country)}

    Please ensure that your translation adheres to the guidelines provided.

    Output ONLY the translation. Do not provide any explanations or text apart from the translation.

    <SOURCE_TEXT>
    """
    for element in elements:
        translation_prompt += f"{element[1]} "
    
    translation_prompt += "</SOURCE_TEXT>"
    return translation_prompt

def get_error_categories(target_lang, country):
    return f"""
### Error Categories:
1. **Accuracy Issues**: 
   - Any duplication, mistranslation, omission, or untranslated text errors that make the translation incorrect.
2. **Language Quality and Punctuation**:
   - Any grammar, spelling, or punctuation errors that make the translation hard to read. Ensure that all words are spelled correctly according to the conventions of {target_lang} and that punctuation conforms to {target_lang} norms.
3. **Style Consistency**:
   - Any mismatch in the style (e.g., news article, court case, literature, casual conversation, financial report, technology essay, etc.) that makes the translation seem unnatural or inappropriate. For example, if the source text is a news article, the translation should read like a news article in {target_lang}. Otherwise, point this out as an error.
4. **Cultural and Regional Appropriateness**:
   - Any mismatch in writing or speaking conventions specific to {country}. For example, in trditional Chinese with zh-TW locale, "disabled" should be translated as "殘障人士" instead of "殘疾人".
    """

# Function to define the error report format for reviewing translations
def get_error_report_format(source_lang, target_lang):
    return f"""
### Error Report Format:
[
  {{
    "original": "string",
    "translated": "string",
    "errors": "string"
  }},
  ...
]

### Attribute Explanations:

- **original**: The sentence in {source_lang}.
- **translated**: The translation of the original sentence in {target_lang}.
- **errors**: The errors identified by experts in the translated sentence. Point out the errors only; do not provide corrected sentences. Each description should be less than 30 words.
"""

# Usage example in a larger prompt
def create_error_report_prompt(source_lang, target_lang, country, elements, translated_text):
    error_report_prompt = f"""
    Your task is to carefully read the following source text within <SOURCE_TEXT></SOURCE_TEXT> and its translation within <TRANSLATION></TRANSLATION> from {source_lang} to {target_lang}. Identify and point out errors in each problematic sentence in an error report.

    Use the following error categories to guide your review:

    ### Error Categories:
    {get_error_categories(target_lang, country)}

    After identifying the errors, document them in the format specified below:

    ### Error Report Format:
    {get_error_report_format(source_lang, target_lang)}

    Output only the error report in the specified format and nothing else.

    <SOURCE_TEXT>
    """
    # Append the source text elements to the prompt
    for element in elements:
        error_report_prompt += f"{element[1]} "

    error_report_prompt += f"</SOURCE_TEXT>\n\n<TRANSLATION>\n{translated_text}\n</TRANSLATION>"

    return error_report_prompt

def create_improvement_prompt(source_lang, target_lang, country, translated_text, error_report):
    improvement_prompt = f"""
    A translation in {target_lang} is provided below within <TRANSLATION></TRANSLATION>. An error report on the translation is provided below within <ERRORS></ERRORS>. The error report follows this format:

    ### Error Report Format:
    {get_error_report_format(source_lang, target_lang)}

    Your task is to:

    - Identify all the problematic translated sentences in the translation provided within <TRANSLATION></TRANSLATION>.
    - Correct the "errors" in these sentences.
    - Provide an improved translation that ensures the text is free from these errors.
    - Maintain the same style: For example, if the source text is a news article, the translation should read like a news article in {target_lang}.
    - Use the appropriate writing or speaking conventions for {country}: For example, in 繁體中文, "disabled" should be "殘障人士" instead of "殘疾人".

    ### Translation Guidelines:
    {get_translation_guidelines(target_lang, country)}

    Output only the improved translation and nothing else.

    <TRANSLATION>
    {translated_text}
    </TRANSLATION>

    <ERRORS>
    {error_report}
    </ERRORS>
    """

    return improvement_prompt

# Array of objects containing language-specific additional notes
notes = [
    {"target_lang": "English", "additional_note": "and adhere to the Chicago Manual of Style"},
    {"target_lang": "French", "additional_note": "and follow the guidelines of the Académie française"},
    {"target_lang": "Spanish", "additional_note": "and adhere to the rules of the Real Academia Española"},
    {"target_lang": "German", "additional_note": "and follow the Duden guidelines for spelling and grammar"},
    {"target_lang": "Italian", "additional_note": "and follow the Accademia della Crusca recommendations"},
    {"target_lang": "Portuguese", "additional_note": "and adhere to the guidelines of the Novo Acordo Ortográfico"},
    {"target_lang": "Russian", "additional_note": "and follow the rules of the Russian Academy of Sciences"},
    {"target_lang": "Japanese", "additional_note": "and adhere to the standards of the Japanese Language Council"},
    {"target_lang": "Korean", "additional_note": "and follow the guidelines of the National Institute of the Korean Language"},
    {"target_lang": "Arabic", "additional_note": "and adhere to the rules of the Arabic Language Academy"},
    {"target_lang": "Dutch", "additional_note": "and follow the Taalunie guidelines for Dutch spelling and grammar"},
    {"target_lang": "Swedish", "additional_note": "and adhere to the guidelines of Svenska Akademiens ordlista"},
    {"target_lang": "Danish", "additional_note": "and follow the guidelines of Dansk Sprognævn"},
    {"target_lang": "Norwegian", "additional_note": "and adhere to the rules of Språkrådet for Bokmål and Nynorsk"},
    {"target_lang": "Finnish", "additional_note": "and follow the guidelines of the Institute for the Languages of Finland"},
    {"target_lang": "Polish", "additional_note": "and adhere to the Polish Language Council's recommendations"},
    {"target_lang": "Turkish", "additional_note": "and follow the guidelines of the Turkish Language Association"},
    {"target_lang": "Greek", "additional_note": "and adhere to the guidelines of the Centre for the Greek Language"},
    {"target_lang": "Hindi", "additional_note": "and follow the guidelines of the Central Hindi Directorate"},
    {"target_lang": "Bengali", "additional_note": "and adhere to the rules of the Bangla Academy"},
    {"target_lang": "Urdu", "additional_note": "and follow the guidelines of the National Language Authority of Pakistan"},
    {"target_lang": "Persian", "additional_note": "and adhere to the guidelines of the Academy of Persian Language and Literature"},
]

def get_additional_note_for_lang(target_lang):
    note = next((item["additional_note"] for item in notes if item["target_lang"] == target_lang), "")
    return note

def create_natural_translation_prompt(target_lang, country, improved_translation, writer):
    additional_note = get_additional_note_for_lang(target_lang)
    natural_translation_prompt = f"""
    You are a skilled writer and editor at a translation agency, fluent in {target_lang}. Your task is to:

    Refine the following translation by infusing the style of {writer}, a renowned 20th-century {target_lang} writer from {country}. The text needs to be in {country} locale {additional_note}.

    Provide only the edited translation, maintaining the essence and tone of the original text, but with the distinct stylistic nuances of {writer}. 

    <TRANSLATION>
    {improved_translation}
    </TRANSLATION>
    """
    print("Generated Natural Translation Prompt:\n", natural_translation_prompt)  # Print the prompt for verification
    return natural_translation_prompt

def create_error_free_translation_prompt(target_lang, country, natural_translation, writer):
    additional_note = get_additional_note_for_lang(target_lang)
    error_free_translation_prompt = f"""
    The following translation has been crafted in the style of {writer}, a famous 20th-century {target_lang} writer from {country}. Your task is to:

    Correct any grammatical, syntactical, or contextual errors while preserving the distinctive style of {writer}. The text needs to be in {country} locale {additional_note}.

    Provide only the edited, error-free translation, ensuring the original tone and style remain intact.

    <TRANSLATION>
    {natural_translation}
    </TRANSLATION>
    """
    print("Generated Error-Free Translation Prompt:\n", error_free_translation_prompt)  # Print the prompt for verification
    return error_free_translation_prompt
