from docx import Document
from docx.shared import Inches

# Load the .docx template
doc = Document('docs/template.docx')

# Define the text to be replaced with an image
target_text = "Insert Image Here"

# Loop through the paragraphs to find and replace the target text
for paragraph in doc.paragraphs:
    if target_text in paragraph.text:
        # Split the paragraph at the target text
        before, target, after = paragraph.text.partition(target_text)
        
        # Clear the paragraph
        paragraph.clear()
        
        # Add the text before the target
        if before:
            paragraph.add_run(before)
        
        # Add the image in place of the target text
        paragraph.add_run().add_picture('images/image_to_insert.png', width=Inches(2))  # Adjust width as needed
        
        # Add the text after the target
        if after:
            paragraph.add_run(after)

# Save the modified document
doc.save('docs/modified_template.docx')
