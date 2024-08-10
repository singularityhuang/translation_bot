from docx import Document
import os
import shutil

def chunk_docx_by_paragraphs(input_file, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    else:
        shutil.rmtree(output_dir)  # Clear the directory if it exists
        os.makedirs(output_dir)

    # Load the input DOCX file
    doc = Document(input_file)

    # Initialize variables to keep track of chunks
    chunk_index = 0
    output_files = []

    # Iterate through paragraphs in the document
    for paragraph in doc.paragraphs:
        # Create a new Document for each paragraph
        new_doc = Document()
        new_doc.add_paragraph(paragraph.text)

        # Save the document with one paragraph
        output_file = os.path.join(output_dir, f"chunk_{chunk_index}.docx")
        new_doc.save(output_file)
        output_files.append(output_file)
        print(f"Saved chunk {chunk_index} successfully.")
        chunk_index += 1

    return output_files

def combine_docx_files(file_list, output_file):
    # Create a new empty Document
    master_doc = Document()

    # Iterate through the list of files and append each one to the master document
    for file in file_list:
        temp_doc = Document(file)
        for paragraph in temp_doc.paragraphs:
            master_doc.add_paragraph(paragraph.text)

    # Save the combined document
    master_doc.save(output_file)
    print(f"Combined document saved as {output_file}")

def main():
    input_file = 'large_document.docx'
    output_dir = 'original_chunks'
    combined_output_file = 'combined.docx'

    # Step 1: Chunk the document by paragraphs, where each chunk is a single paragraph
    chunked_files = chunk_docx_by_paragraphs(input_file, output_dir)

    # Step 2: Combine the chunked files into one document (if needed)
    combine_docx_files(chunked_files, combined_output_file)

if __name__ == "__main__":
    main()
