from docx import Document

# Load the original document
doc = Document('your_input_file.docx')

# Initialize a counter for the sequential numbers
counter = 1

# Function to add sequential number to text elements
def add_sequential_number_to_paragraph(paragraph, counter):
    if paragraph.runs:
        for run in paragraph.runs:
            run.text += f" {counter}"
        counter += 1
    else:
        paragraph.add_run(f" {counter}")
        counter += 1
    return counter

# Iterate through all paragraphs in the document body
for paragraph in doc.paragraphs:
    counter = add_sequential_number_to_paragraph(paragraph, counter)

# Iterate through all tables in the document
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # Iterate through each paragraph in the cell
            for paragraph in cell.paragraphs:
                counter = add_sequential_number_to_paragraph(paragraph, counter)

# Save the modified document
doc.save('output_with_sequential_numbers.docx')

print("Added sequential numbers to each paragraph and table cell in the document.")
