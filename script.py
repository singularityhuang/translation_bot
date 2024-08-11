import os
import shutil
import requests
from dotenv import load_dotenv
import docx
from docx import Document
import argparse
import re
from translate_prompt import extract_text_with_structure, create_translation_prompt, create_error_report_prompt, create_improvement_prompt, create_natural_translation_prompt, create_error_free_translation_prompt

# Load environment variables from the .env file
load_dotenv()

# Retrieve the API key from the environment variables
api_key = os.getenv("GEMINI_API_KEY")

def translate_texts(source_lang, target_lang, country, writer):
    # Initialize variables
    token_counts = []
    total_tokens_used = 0
    # Define your input parameters
    source_folder = "original_chunks"
    final_translation_file = "final_translation.docx"

    # Create a new document for the final merged content
    final_doc = Document()

    # List and sort all .docx files in the source folder numerically
    files = [f for f in os.listdir(source_folder) if f.startswith("chunk_") and f.endswith(".docx")]
    files.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))

    # Iterate over all .docx files in the source folder in sorted order
    for i, filename in enumerate(files):
        source_docx_file = os.path.join(source_folder, filename)
        print(f"Processing {source_docx_file}")

        # Extract text with structure
        elements = extract_text_with_structure(source_docx_file)

        # Check if the extracted text is empty or contains only line breaks
        text_content = "".join(element[1] for element in elements).strip()
        if not text_content:
            print(f"Original content for {source_docx_file} contains only line breaks or is empty.")
            append_original_to_docx(final_doc, source_docx_file)
            continue

        # Create the translation prompt
        translation_prompt = create_translation_prompt(source_lang, target_lang, country, elements)

        # Setup the URL and headers for the API request
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={api_key}"
        headers = {
            "Content-Type": "application/json"
        }

        def call_gemini_api(prompt):
            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "text": prompt
                            }
                        ]
                    }
                ]
            }
            response = requests.post(url, headers=headers, json=payload)
            if response.status_code == 200:
                response_data = response.json()
                print("Response JSON:", response_data)
                token_count = response_data.get('usageMetadata', {}).get('totalTokenCount', 0)
                return response_data, token_count
            else:
                print(f"Error: {response.status_code} - {response.text}")
                return None, 0

        def process_translation(prompt):
            """Helper function to handle translation and retry if necessary."""
            translation_response, tokens = call_gemini_api(prompt)
            token_counts.append(tokens)  # Accumulate tokens

            if translation_response and 'candidates' in translation_response and 'content' in translation_response['candidates'][0]:
                return translation_response, tokens
            else:
                print("Retrying due to missing 'content' in response...")
                # Retry once
                translation_response, tokens = call_gemini_api(prompt)
                token_counts.append(tokens)  # Accumulate tokens
                if translation_response and 'candidates' in translation_response and 'content' in translation_response['candidates'][0]:
                    return translation_response, tokens
                else:
                    print("Translation failed after retry.")
                    return None, tokens

        # Call the Gemini API for translation
        translation_response, tokens = process_translation(translation_prompt)
        if translation_response:
            # Extract the translated text from the response
            translated_text = translation_response['candidates'][0]['content']['parts'][0]['text']
            translated_text = re.sub(r'<[^>]+>', '', translated_text).strip()

            # Continue with further processing (error report, improvement, natural translation, etc.)
            error_report_prompt = create_error_report_prompt(source_lang, target_lang, country, elements, translated_text)
            error_report_response, tokens = process_translation(error_report_prompt)
            if error_report_response:
                error_report_text = error_report_response['candidates'][0]['content']['parts'][0]['text']

                improvement_prompt = create_improvement_prompt(source_lang, target_lang, country, translated_text, error_report_text)
                improvement_response, tokens = process_translation(improvement_prompt)
                if improvement_response:
                    improvement_text = improvement_response['candidates'][0]['content']['parts'][0]['text']
                    improvement_text = re.sub(r'<[^>]+>', '', improvement_text).strip()

                    natural_translation_prompt = create_natural_translation_prompt(target_lang, country, improvement_text, writer)
                    natural_translation_response, tokens = process_translation(natural_translation_prompt)
                    if natural_translation_response:
                        natural_translation_text = natural_translation_response['candidates'][0]['content']['parts'][0]['text'].replace('\n\n','')
                        natural_translation_text = re.sub(r'<[^>]+>', '', natural_translation_text).strip()

                        error_free_translation_prompt = create_error_free_translation_prompt(target_lang, country, natural_translation_text, writer)
                        error_free_translation_response, tokens = process_translation(error_free_translation_prompt)
                        if error_free_translation_response:
                            error_free_translation_text = error_free_translation_response['candidates'][0]['content']['parts'][0]['text'].replace('\n\n','')
                            error_free_translation_text = re.sub(r'<[^>]+>', '', error_free_translation_text).strip()

                            # Append the final error-free translation to the final document
                            append_translation_to_docx(final_doc, error_free_translation_text, elements)
                            continue  # Move to the next chunk

        # If we reached here, it means translation or any subsequent step failed
        print(f"Appending original content for {source_docx_file} due to translation failure or missing content.")
        append_original_to_docx(final_doc, source_docx_file)

    # Calculate and print the total number of tokens used
    total_tokens_used = sum(token_counts)
    print(f"Total tokens used: {total_tokens_used}")

    # Save the final merged document
    final_doc.save(final_translation_file)
    print(f"All translations saved to {final_translation_file}")




def save_text_to_docx(source_docx_file, text, elements, destination_docx_file, is_plain_text=False):
    # Replace \n\n with '' within the text to keep paragraphs intact
    cleaned_text = text.replace('\n\n', '')

    # Split text by single newline for processing
    paragraphs = cleaned_text.split('\n')
    
    # Create a new docx document
    doc = Document()
    
    # If saving plain text, directly add the paragraphs
    if is_plain_text:
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    else:
        # Map text to the original elements structure
        for element, para in zip(elements, paragraphs):
            if element[0] == 'paragraph':
                new_para = doc.add_paragraph(style=element[2])
                new_run = new_para.add_run(para)
                copy_run_formatting(new_run, element[3][0] if element[3] else None)
            elif element[0] == 'cell':
                new_table = doc.add_table(rows=1, cols=1)
                cell = new_table.cell(0, 0)
                cell.text = para
                for run in element[3]:
                    new_run = cell.paragraphs[0].add_run(run.text)
                    copy_run_formatting(new_run, run)

    # Save the document
    doc.save(destination_docx_file)
    print(f"Text saved to {destination_docx_file}")

def copy_run_formatting(new_run, original_run):
    if original_run:
        new_run.bold = original_run.bold
        new_run.italic = original_run.italic
        new_run.underline = original_run.underline
        new_run.font.size = original_run.font.size
        new_run.font.color.rgb = original_run.font.color.rgb

def append_translation_to_docx(merged_doc, translated_text, elements):
    translated_paragraphs = translated_text.split('\n')

    # Map translated text to the original elements structure
    for element, translated_para in zip(elements, translated_paragraphs):
        if element[0] == 'paragraph':
            new_para = merged_doc.add_paragraph(style=element[2])
            new_run = new_para.add_run(translated_para)
            copy_run_formatting(new_run, element[3][0] if element[3] else None)
        elif element[0] == 'cell':
            new_table = merged_doc.add_table(rows=1, cols=1)
            cell = new_table.cell(0, 0)
            cell.text = translated_para
            for run in element[3]:
                new_run = cell.paragraphs[0].add_run(run.text)
                copy_run_formatting(new_run, run)

def append_original_to_docx(merged_doc, source_docx_file):
    original_doc = Document(source_docx_file)
    for paragraph in original_doc.paragraphs:
        new_para = merged_doc.add_paragraph(style=paragraph.style)
        for run in paragraph.runs:
            new_run = new_para.add_run(run.text)
            copy_run_formatting(new_run, run)


def merge_error_free_chunks(final_doc, original_folder, error_free_folder):
    # List all .docx files in the original folder that match the pattern chunk_{i}.docx
    files = [f for f in os.listdir(original_folder) if f.startswith("chunk_") and f.endswith(".docx")]

    # Sort files numerically by extracting the number after 'chunk_' and before '.docx'
    files.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))

    # Iterate over sorted files and merge the corresponding documents
    for filename in files:
        original_docx_file = os.path.join(original_folder, filename)
        error_free_docx_file = os.path.join(error_free_folder, filename)

        if os.path.exists(error_free_docx_file):
            doc_to_append = Document(error_free_docx_file)
            append_doc_to_another(final_doc, doc_to_append)
        else:
            doc_to_append = Document(original_docx_file)
            append_doc_to_another(final_doc, doc_to_append)


def append_doc_to_another(doc, doc_to_append):
    for element in iter_block_items(doc_to_append):
        if isinstance(element, docx.text.paragraph.Paragraph):
            new_para = doc.add_paragraph(style=element.style)
            for run in element.runs:
                new_run = new_para.add_run(run.text)
                copy_run_formatting(new_run, run)
        elif isinstance(element, docx.table.Table):
            new_table = doc.add_table(rows=0, cols=len(element.columns))
            for row in element.rows:
                new_row = new_table.add_row()
                for i, cell in enumerate(row.cells):
                    new_cell = new_row.cells[i]
                    for paragraph in cell.paragraphs:
                        new_para = new_cell.add_paragraph(style=paragraph.style)
                        for run in paragraph.runs:
                            new_run = new_para.add_run(run.text)
                            copy_run_formatting(new_run, run)

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._element
    else:
        raise ValueError("Unknown parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.CT_Tbl):
            yield docx.table.Table(child, parent)

if __name__ == "__main__":
    # Set up argument parser
    parser = argparse.ArgumentParser(description="Translation script parameters")
    parser.add_argument('--source_lang', type=str, default="English", help="Source language")
    parser.add_argument('--target_lang', type=str, default="Traditional Chinese", help="Target language")
    parser.add_argument('--country', type=str, default="Taiwan", help="Country associated with the translation")
    parser.add_argument('--writer', type=str, default="龍應台", help="Writer for natural translation context")
    
    args = parser.parse_args()

    translate_texts(args.source_lang, args.target_lang, args.country, args.writer)
