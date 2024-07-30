from docx import Document
import os
import shutil
from copy import deepcopy

def append_element(new_doc, element, element_index):
    """
    Helper function to append an element along with its formatting.
    """
    try:
        new_element = deepcopy(element)
        new_doc._element.body.append(new_element)
    except Exception as e:
        print(f"Error appending element at index {element_index}: {e}")

def chunk_docx_by_elements(input_file, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    else:
        shutil.rmtree(output_dir)  # Clear the directory if it exists
        os.makedirs(output_dir)

    # Load the input DOCX file
    doc = Document(input_file)

    # Initialize variables to keep track of elements
    chunk_index = 0
    output_files = []

    # Create a new Document for each chunk
    new_doc = Document()
    element_index = 0
    for element in doc.element.body:
        append_element(new_doc, element, element_index)

        # Save and start a new document if a table or paragraph is encountered
        if element.tag.endswith(('tbl', 'p')):
            output_file = os.path.join(output_dir, f"chunk_{chunk_index}.docx")
            try:
                new_doc.save(output_file)
                output_files.append(output_file)
                print(f"Saved chunk {chunk_index} successfully.")
            except Exception as e:
                print(f"Error saving chunk {chunk_index}: {e}")
            chunk_index += 1
            new_doc = Document()

        element_index += 1

    # Save the last document if it contains content
    if len(new_doc.element.body):
        output_file = os.path.join(output_dir, f"chunk_{chunk_index}.docx")
        try:
            new_doc.save(output_file)
            output_files.append(output_file)
            print(f"Saved final chunk {chunk_index} successfully.")
        except Exception as e:
            print(f"Error saving final chunk {chunk_index}: {e}")

    return output_files

def combine_docx_files(file_list, output_file):
    # Create a new empty Document
    master_doc = Document()

    # Iterate through the list of files and append each one to the master document
    file_index = 0
    for file in file_list:
        temp_doc = Document(file)
        element_index = 0
        for element in temp_doc.element.body:
            append_element(master_doc, element, f"file {file_index}, element {element_index}")
            element_index += 1
        file_index += 1

    # Save the combined document
    try:
        master_doc.save(output_file)
        print(f"Combined document saved as {output_file}")
    except Exception as e:
        print(f"Error saving combined document: {e}")

def main():
    input_file = 'large_document.docx'
    output_dir = 'original_chunks'
    combined_output_file = 'combined.docx'

    # Step 1: Chunk the document by elements
    chunked_files = chunk_docx_by_elements(input_file, output_dir)

    # Step 2: Combine the chunked files into one document
    combine_docx_files(chunked_files, combined_output_file)

if __name__ == "__main__":
    main()
