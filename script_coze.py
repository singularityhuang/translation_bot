import docx
import math
import os
import requests
import json
import time

def load_docx(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def split_text(text, chunk_size):
    words = text.split()
    num_chunks = math.ceil(len(words) / chunk_size)
    return [' '.join(words[i*chunk_size:(i+1)*chunk_size]) for i in range(num_chunks)]

def save_chunks(chunks):
    if not os.path.exists('chunks'):
        os.makedirs('chunks')
    for i, chunk in enumerate(chunks):
        doc = docx.Document()
        doc.add_paragraph(chunk)
        doc.save(f'chunks/chunk_{i+1}.docx')

def combine_chunks(translated_chunks):
    doc = docx.Document()
    for chunk in translated_chunks:
        doc.add_paragraph(chunk)
    doc.save('translated_full_document.docx')

def translate_text(chunk, access_token, bot_id, conversation_id, user, retry_count=3):
    url = 'https://api.coze.com/open_api/v2/chat'
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json',
        'Accept': '*/*',
        'Host': 'api.coze.com',
        'Connection': 'keep-alive'
    }
    data = {
        "conversation_id": conversation_id,
        "bot_id": bot_id,
        "user": user,
        "query": f"請將下面的文章翻譯成中文：{chunk}",
        "stream": False
    }
    for attempt in range(retry_count):
        try:
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=300)
            if response.status_code == 200:
                response_data = response.json()
                if 'messages' in response_data:
                    for message in response_data['messages']:
                        if "最終版的翻譯" in message['content']:
                            return message['content']
                else:
                    print("Error: 'messages' key not found in the response")
                    print(response_data)
                    return ""
            else:
                print(f"Error: Received status code {response.status_code}")
                print(response.text)
                if response.status_code in [500, 502, 503, 504]:  # Retry only for specific error codes
                    wait_time = 2 ** attempt
                    print(f"Retrying in {wait_time} seconds...")
                    time.sleep(wait_time)
                else:
                    break
        except requests.exceptions.RequestException as e:
            print(f"Request failed: {e}")
            wait_time = 2 ** attempt
            print(f"Retrying in {wait_time} seconds...")
            time.sleep(wait_time)
    
    return ""

if __name__ == "__main__":
    # Configuration
    access_token = 'pat_ZbPl7FjFj0IvxvGUHWzpgVB4BwY7IVoxPpOS558tp3J5VQlkyFOmPHYkPhsGjXXi'
    bot_id = '7391097348502011909'
    conversation_id = '123'
    user = '29032201862555'

    # Load the document
    text = load_docx('large_document.docx')

    # Define the chunk size (test with 1000 words)
    chunk_size = 100

    # Split the text into chunks
    chunks = split_text(text, chunk_size)

    # Save the chunks as separate docx files
    save_chunks(chunks)

    print(f"Document split into {len(chunks)} chunks and saved in 'chunks' directory.")

    # Translate each chunk
    translated_chunks = []
    for i, chunk in enumerate(chunks):
        print(f"Translating chunk {i+1}/{len(chunks)}")  # Log the start of each chunk translation
        translated_text = translate_text(chunk, access_token, bot_id, conversation_id, user)
        translated_chunks.append(translated_text)
        print(f"Translated chunk {i+1}/{len(chunks)}")
        print(f"{translated_text}")

    # Combine translated chunks into a single document
    combine_chunks(translated_chunks)

    print("Translated chunks combined into 'translated_full_document.docx'.")
